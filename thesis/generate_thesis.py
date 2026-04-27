from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import re
import textwrap


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "thesis"
ASSETS = OUT / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Smart_Campus_Navigation_Thesis_Draft.docx"

TITLE = "Design and Implementation of a Smart Campus Navigation System"


OUTLINE = """
Chapter 1 Introduction
1.1 Research Background and Significance
1.1.1 Background of Smart Campus Development
1.1.2 Demand for Campus Navigation Services
1.1.3 Influence of Weather on Campus Travel
1.2 Research Status at Home and Abroad
1.2.1 Research Status of Map Navigation and Location-Based Services
1.2.2 Research Status of Campus Navigation Systems
1.2.3 Research Status of Weather-Aware Travel Services
1.3 Main Work of This Thesis
1.3.1 Campus POI Retrieval and Place Detail Display
1.3.2 Route Planning and Map Interaction Based on AMap (Gaode Map)
1.3.3 Weather Display, Travel Suggestions, and Intelligent Assistant Q&A
1.4 Organization of the Thesis
Chapter 2 Development Tools and Related Technologies
2.1 Development Environment
2.1.1 Java 21 and Maven
2.1.2 IntelliJ IDEA Development Tool
2.1.3 MySQL Database Environment
2.2 Backend Development Technologies
2.2.1 Spring Boot 3.5.13
2.2.2 Spring MVC Request Processing Mechanism
2.2.3 MyBatis Persistence Framework
2.2.4 Unified Response and Exception Handling Mechanism
2.3 Frontend Development Technologies
2.3.1 Thymeleaf Template Engine
2.3.2 HTML, CSS, and Native JavaScript
2.3.3 Map Interaction and Sidebar State Management
2.4 Third-Party Service Technologies
2.4.1 AMap JavaScript API
2.4.2 AMap Web Service Route Service
2.4.3 QWeather API
2.4.4 Backend Proxy Interface for the Intelligent Assistant
2.5 System Testing Methods
2.5.1 Functional Testing
2.5.2 Interface Testing
2.5.3 Exception Testing
Chapter 3 System Analysis and Design
3.1 System Feasibility Analysis
3.1.1 Technical Feasibility
3.1.2 Economic Feasibility
3.1.3 Operational Feasibility
3.2 System Boundary Description
3.2.1 User System Boundary
3.2.2 Data Storage Boundary
3.2.3 Functional Extension Boundary
3.3 Requirement Analysis
3.3.1 Functional Requirement Analysis
3.3.2 Non-Functional Requirement Analysis
3.4 Overall System Architecture Design
3.4.1 Presentation Layer Design
3.4.2 Controller Layer Design
3.4.3 Service Layer Design
3.4.4 Persistence Layer Design
3.4.5 Third-Party Service Interaction Design
3.5 System Functional Structure Design
3.5.1 POI Retrieval and Place Detail Module
3.5.2 Route Planning Module
3.5.3 Weather Display and Travel Suggestion Module
3.5.4 Saved Places Module
3.5.5 Route History Records Module
3.5.6 Intelligent Assistant Module
3.6 Database Design
3.6.1 Conceptual Structure Design
3.6.2 Logical Structure Design
3.6.3 Description of Main Data Tables
3.6.4 Description of Data Table Relationships
3.7 Core Business Process Design
3.7.1 POI Query Process Design
3.7.2 Route Planning Process Design
3.7.3 Weather-Aware Travel Suggestion Process Design
3.7.4 Intelligent Assistant Q&A Process Design
3.8 System Interface Design
3.8.1 POI-Related Interfaces
3.8.2 Route Planning Interfaces
3.8.3 Weather and Suggestion Interfaces
3.8.4 Saved Places and Route History Records Interfaces
3.8.5 Intelligent Assistant Interface
3.8.6 Main System Interface Description Table
3.9 Chapter Summary
Chapter 4 System Implementation
4.1 Engineering Structure Description
4.1.1 Project Directory Structure
4.1.2 Main Package Responsibilities
4.1.3 Resource File Organization
4.2 Home Page and Map Initialization Implementation
4.2.1 Home Page Layout Design
4.2.2 Map Container Initialization
4.2.3 Campus Center Point and Map Component Loading
4.2.4 Home Page Running Effect
4.3 POI Retrieval and Place Detail Implementation
4.3.1 Search Box and Result List Interaction
4.3.2 POI Query Request Processing
4.3.3 Place Detail Data Rendering
4.3.4 Map Marker and Location Display
4.3.5 POI Retrieval Running Effect
4.4 Route Planning Implementation
4.4.1 Start Point, Waypoint, and Destination Selection
4.4.2 Route Planning Request Construction
4.4.3 AMap Route Service Invocation
4.4.4 Route Result Parsing and Frontend Map Rendering
4.4.5 Route Planning Running Effect
4.5 Weather Display and Weather Risk Prompt Implementation
4.5.1 Weather Information Display Area Design
4.5.2 QWeather API Request Processing
4.5.3 Weather Risk Prompt Logic
4.5.4 Weather Information Exception Fallback
4.5.5 Weather Display Running Effect
4.6 Travel Suggestion Implementation
4.6.1 Scenario-Based Suggestion Rules
4.6.2 Combination of Route and Weather Information
4.6.3 Assisted Processing of Sheltered Path Data
4.6.4 Frontend Rendering of Suggestion Results
4.6.5 Travel Suggestion Running Effect
4.7 Saved Places Implementation
4.7.1 Adding Saved Places
4.7.2 Saved Places List Display
4.7.3 Saved Place Name Editing
4.7.4 Saved Place Deletion and Data Persistence
4.7.5 Saved Places Running Effect
4.8 Route History Records Implementation
4.8.1 Route History Record Storage
4.8.2 Route History Records List Display
4.8.3 Route History Record Reuse
4.8.4 Route History Record Deletion and Cleanup
4.8.5 Route History Records Running Effect
4.9 Intelligent Assistant Implementation
4.9.1 Q&A Entry and Message Display
4.9.2 Campus Scenario Context Construction
4.9.3 Backend Proxy Interface Invocation
4.9.4 Answer Display and Exception Prompt
4.9.5 Intelligent Assistant Running Effect
4.10 Frontend and Backend Exception Handling
4.10.1 Backend Unified Response Structure
4.10.2 Global Exception Handling
4.10.3 Frontend Request Timeout Handling
4.10.4 User Prompt Message Display
4.11 Key Problems and Solutions
4.11.1 Handling Third-Party Weather Interface Failures
4.11.2 Synchronization Between Map Interaction and Sidebar State
4.11.3 Selection and Search Result Refill of Start Point, Waypoint, and Destination
4.11.4 Migration of Saved Places from localStorage to MySQL Persistence
4.11.5 Backend Encapsulation of External API Keys
4.12 Chapter Summary
Chapter 5 System Testing
5.1 Testing Environment
5.1.1 Hardware Environment
5.1.2 Software Environment
5.1.3 Database Environment
5.2 Testing Methods
5.2.1 Functional Testing
5.2.2 Interface Testing
5.2.3 Exception Testing
5.3 Functional Testing
5.3.1 POI Query and Place Detail Testing
5.3.2 Route Planning Testing
5.3.3 Weather Display and Travel Suggestion Testing
5.3.4 Saved Places Testing
5.3.5 Route History Records Testing
5.3.6 Intelligent Assistant Testing
5.4 Exception Testing
5.4.1 Search Exception Testing
5.4.2 Route Parameter Exception Testing
5.4.3 Third-Party Interface Exception Testing
5.4.4 No Database Query Result Testing
5.5 Interface Testing
5.5.1 POI Interface Testing
5.5.2 Route Planning Interface Testing
5.5.3 Weather and Suggestion Interface Testing
5.5.4 Saved Places and Route History Records Interface Testing
5.5.5 Intelligent Assistant Interface Testing
5.6 Test Result Analysis
5.7 Chapter Summary
Chapter 6 Conclusion and Future Work
6.1 Conclusion
6.1.1 System Completion Summary
6.1.2 Main Implementation Results
6.1.3 System Limitations
6.2 Future Work
6.2.1 Continuous Improvement of Campus POI Data
6.2.2 Mobile Adaptation Optimization
6.2.3 Extension to Indoor Navigation and More Fine-Grained Path Services
""".strip().splitlines()


FIGURES = {
    "Chapter 1 Introduction": ("Figure 1.1 System Application Scenario Diagram", "scenario"),
    "Chapter 2 Development Tools and Related Technologies": ("Figure 2.1 System Technology Stack Diagram", "tech"),
    "3.4 Overall System Architecture Design": ("Figure 3.1 Overall System Architecture Diagram", "architecture"),
    "3.5 System Functional Structure Design": ("Figure 3.2 System Functional Structure Diagram", "function"),
    "3.7.1 POI Query Process Design": ("Figure 3.3 POI Query Flowchart", "poi-flow"),
    "3.7.2 Route Planning Process Design": ("Figure 3.4 Route Planning Flowchart", "route-flow"),
    "3.7.3 Weather-Aware Travel Suggestion Process Design": ("Figure 3.5 Weather-Aware Travel Suggestion Flowchart", "weather-flow"),
    "3.6.1 Conceptual Structure Design": ("Figure 3.6 Database ER Diagram", "er"),
    "4.1.1 Project Directory Structure": ("Figure 4.1 Project Directory Structure Diagram", "directory"),
    "4.2.4 Home Page Running Effect": ("Figure 4.2 System Home Page Screenshot", "home screenshot placeholder"),
    "4.3.5 POI Retrieval Running Effect": ("Figure 4.3 POI Search and Place Detail Screenshot", "poi screenshot placeholder"),
    "4.4.5 Route Planning Running Effect": ("Figure 4.4 Route Planning Result Screenshot", "route screenshot placeholder"),
    "4.5.5 Weather Display Running Effect": ("Figure 4.5 Weather and Suggestion Screenshot", "weather screenshot placeholder"),
    "4.7.5 Saved Places Running Effect": ("Figure 4.6 Saved Places Management Screenshot", "saved places screenshot placeholder"),
    "4.8.5 Route History Records Running Effect": ("Figure 4.7 Route History Records Screenshot", "route history screenshot placeholder"),
    "4.9.5 Intelligent Assistant Running Effect": ("Figure 4.8 Intelligent Assistant Screenshot", "intelligent assistant screenshot placeholder"),
}


def font(size=18, bold=False):
    for name in (("arialbd.ttf" if bold else "arial.ttf"), "Calibri.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap(draw, text, xy, w, f, fill=(30, 40, 50)):
    x, y = xy
    for line in textwrap.wrap(text, width=max(12, w // 10)):
        draw.text((x, y), line, font=f, fill=fill)
        y += 25


def make_figure(caption, kind):
    safe = re.sub(r"[^A-Za-z0-9]+", "_", caption).strip("_").lower() + ".png"
    path = ASSETS / safe
    img = Image.new("RGB", (1200, 720), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1200, 72], fill=(25, 74, 110))
    d.text((36, 22), caption, font=font(24, True), fill="white")
    if "screenshot placeholder" in kind:
        d.rounded_rectangle([90, 115, 1110, 640], radius=18, fill=(247, 249, 251), outline=(90, 110, 130), width=2)
        d.rectangle([90, 115, 1110, 175], fill=(35, 85, 120))
        d.text((120, 133), "Verified screenshot placeholder", font=font(20, True), fill="white")
        d.rounded_rectangle([130, 220, 390, 570], radius=10, fill=(255, 255, 255), outline=(140, 155, 170), width=2)
        d.rounded_rectangle([430, 220, 1060, 570], radius=10, fill=(232, 242, 236), outline=(70, 120, 90), width=2)
        d.text((460, 360), kind.replace(" screenshot placeholder", "").title(), font=font(28, True), fill=(45, 80, 65))
        d.text((185, 610), "Replace this placeholder with an actual verified system screenshot before final submission.", font=font(16), fill=(150, 50, 40))
    elif kind == "directory":
        lines = ["smart-campus-navigation/", "src/main/java/.../controller", "src/main/java/.../service/impl", "src/main/java/.../mapper", "src/main/java/.../entity, dto, vo", "src/main/resources/mapper", "src/main/resources/sql", "src/main/resources/templates/index.html", "src/main/resources/static/css, js, icons"]
        d.rounded_rectangle([150, 115, 1050, 650], radius=12, fill=(247, 249, 252), outline=(95, 111, 128), width=2)
        y = 150
        for line in lines:
            d.text((190, y), line, font=font(19), fill=(20, 35, 50))
            y += 52
    else:
        labels = {
            "scenario": ["Users", "Web System", "MySQL POI Data", "AMap", "QWeather API", "Intelligent Assistant Proxy"],
            "tech": ["Frontend", "Backend", "Persistence", "Third-Party Services"],
            "architecture": ["Presentation Layer", "Controller Layer", "Service Layer", "Persistence Layer", "External APIs"],
            "function": ["POI", "Route", "Weather", "Suggestions", "Saved Places", "Route History", "Intelligent Assistant"],
            "poi-flow": ["Input", "Request", "Controller", "Service", "MySQL", "Render"],
            "route-flow": ["Select Points", "Validate", "AMap Service", "Parse Result", "Render Route"],
            "weather-flow": ["Request Weather", "QWeather API", "Risk Prompt", "Suggestion", "Display"],
            "er": ["poi", "saved_place", "route_history", "covered_path_node"],
        }.get(kind, ["Design Item"])
        x, y = 80, 170
        for i, label in enumerate(labels):
            d.rounded_rectangle([x, y, x + 220, y + 90], radius=12, fill=(244, 248, 252), outline=(45, 86, 130), width=2)
            wrap(d, label, (x + 18, y + 26), 180, font(18))
            if i < len(labels) - 1:
                d.line([x + 220, y + 45, x + 260, y + 45], fill=(70, 80, 90), width=3)
                d.polygon([(x + 260, y + 45), (x + 248, y + 37), (x + 248, y + 53)], fill=(70, 80, 90))
            x += 270
            if x > 950:
                x, y = 180, y + 170
    img.save(path)
    return path


def level_of(title):
    if title.startswith("Chapter "):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\s", title):
        return 3
    return 2


def add_field_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def para(doc, text="", align=None, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0.74 if indent else 0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r)
    return p


def heading(doc, text, lvl):
    p = doc.add_paragraph(style=f"Heading {lvl}")
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(text)
    return p


def body_for(title):
    t = title.lower()
    if title.startswith("Chapter"):
        return []
    if "boundary" in t:
        return ["This section defines the implemented scope of the current Web system. The system does not include login, registration, permission management, a multi-user system, user profiles, model training, deep learning algorithms, vector retrieval, or RAG. Saved Places and Route History Records are stored as shared datasets, so no user table or user_id is designed."]
    if "amap" in t or "route planning" in t:
        return ["Route planning is implemented by validating the selected start point, optional waypoints, and destination, invoking AMap route services, parsing the returned route data, and rendering the route result on the frontend map. The thesis must not describe this function as a self-developed A*, Dijkstra, or route optimization algorithm."]
    if "qweather" in t or "weather" in t:
        return ["The weather-related module obtains current weather data from QWeather API through backend request processing. The system converts weather conditions into practical campus travel prompts and provides exception fallback when the third-party weather service is unavailable."]
    if "intelligent assistant" in t:
        return ["The Intelligent Assistant is described only as backend proxy invocation, campus scenario context construction, answer display, and exception prompt. The project does not train a large model, build a knowledge base system, use vector retrieval, or implement RAG."]
    if "saved places" in t:
        return ["The Saved Places module supports adding, displaying, renaming, and deleting saved campus places. These records are persisted in MySQL as a shared dataset rather than as user-specific favorites."]
    if "route history" in t:
        return ["Route History Records store route-related information such as title, mode, start point, destination, optional waypoint data, distance, and duration. The system supports listing, reusing, deleting, and clearing these shared route history records."]
    if "poi" in t or "place detail" in t:
        return ["The POI module retrieves campus place data from MySQL through MyBatis and displays search results and place details in the frontend sidebar. Place detail rendering includes name, type, description, opening hours, status, and coordinate-based map location display."]
    if "interface" in t:
        return ["Interface paths must be checked against the real Controller classes before final submission. The interface table in this draft uses implemented Controller paths only and should be rechecked after any code change."]
    if "testing" in t or "test" in t:
        return ["Testing focuses on functional behavior, interface responses, and exception scenarios. Test case tables use the fields test ID, test module, test content, input data, expected result, actual result, and pass or fail."]
    if "database" in t or "data table" in t:
        return ["The database design is based on the implemented tables poi, saved_place, route_history, and covered_path_node. The current design does not introduce a user table because the project does not implement a user account system."]
    if "summary" in t:
        return ["This chapter summary reviews the main points of the chapter and connects them with the following chapter. It should remain concise and should not introduce new functions that are not implemented in the project."]
    return ["This section explains the corresponding part of the Smart Campus Navigation Web system according to the implemented project. The writing should stay close to the real technology stack and avoid adding functions that are not present in the codebase."]


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


def set_cell(cell, text, bold=False, size=9):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 26 else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run_font(r, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, data in edges.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        for k, v in data.items():
            el.set(qn(f"w:{k}"), str(v))


def set_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def table(doc, title, headers, rows):
    if len(headers) >= 5:
        wide_table(doc, title, headers, rows)
        return
    p = para(doc, title, WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(10.5)
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    tb.autofit = False
    if len(headers) == 7:
        widths = [1.3, 1.8, 2.5, 2.4, 3.0, 2.2, 1.4]
        size = 7
    elif len(headers) == 5:
        widths = [3.0, 1.5, 3.3, 4.0, 3.0]
        size = 7
    elif len(headers) == 3:
        widths = [2.0, 3.0, 10.0]
        size = 9
    else:
        widths = [15.0 / len(headers)] * len(headers)
        size = 9
    for i, h in enumerate(headers):
        set_width(tb.rows[0].cells[i], widths[i])
        set_cell(tb.rows[0].cells[i], h, True, size)
    for row in rows:
        cells = tb.add_row().cells
        for i, val in enumerate(row):
            set_width(cells[i], widths[i])
            set_cell(cells[i], val, False, size)
    for row in tb.rows:
        for c in row.cells:
            border(c, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    for c in tb.rows[0].cells:
        border(c, top={"val": "single", "sz": "12", "color": "000000"}, bottom={"val": "single", "sz": "8", "color": "000000"}, left={"val": "nil"}, right={"val": "nil"})
    for c in tb.rows[-1].cells:
        border(c, bottom={"val": "single", "sz": "12", "color": "000000"}, left={"val": "nil"}, right={"val": "nil"})
    para(doc, "", indent=False)


def figure(doc, caption, kind):
    path = make_figure(caption, kind)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(14.5))
    c = para(doc, caption, WD_ALIGN_PARAGRAPH.CENTER, False)
    c.runs[0].font.size = Pt(10.5)
    para(doc, "", indent=False)


def table_image(title, headers, rows):
    safe = re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_").lower() + ".png"
    path = ASSETS / safe
    if len(headers) == 7:
        col_w = [150, 210, 300, 300, 360, 260, 170]
    else:
        col_w = [360, 180, 360, 500, 300]
    width = sum(col_w) + 80
    f = font(22)
    fb = font(22, True)

    def cell_lines(text, w):
        return textwrap.wrap(str(text), width=max(8, w // 13)) or [""]

    all_lines = []
    heights = []
    for row in [headers] + rows:
        lines = [cell_lines(v, col_w[i] - 16) for i, v in enumerate(row)]
        all_lines.append(lines)
        heights.append(max(44, max(len(x) for x in lines) * 28 + 18))

    height = sum(heights) + 80
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    x0, y = 40, 36
    d.line([x0, y, width - 40, y], fill="black", width=4)
    for r_idx, lines in enumerate(all_lines):
        x = x0
        for c_idx, cell in enumerate(lines):
            cy = y + 10
            for line in cell:
                d.text((x + 8, cy), line, font=fb if r_idx == 0 else f, fill="black")
                cy += 28
            x += col_w[c_idx]
        y += heights[r_idx]
        if r_idx == 0:
            d.line([x0, y, width - 40, y], fill="black", width=3)
    d.line([x0, y, width - 40, y], fill="black", width=4)
    img.save(path)
    return path


def wide_table(doc, title, headers, rows):
    p = para(doc, title, WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(10.5)
    img = table_image(title, headers, rows)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(img), width=Cm(15.5))
    para(doc, "", indent=False)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    for style_name, size, bold in [("Normal", 12, False), ("Heading 1", 14, True), ("Heading 2", 12, True), ("Heading 3", 12, False)]:
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.first_line_indent = Cm(0 if style_name.startswith("Heading") else 0.74)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("南京信息工程大学本科毕业论文（设计）")
    hr.font.name = "SimSun"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    hr.font.size = Pt(9)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field_page_number(fp)

    para(doc, "Unit Code: 10300", indent=False)
    para(doc, "", indent=False)
    p = para(doc, "Undergraduate Graduation Thesis (Design)", WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(18)
    p.runs[0].bold = True
    para(doc, "", indent=False)
    p = para(doc, "Title: " + TITLE, WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    for item in ["Student Name: [To be filled]", "Student ID: [To be filled]", "Major: [To be filled]", "College: [To be filled]", "Supervisor: [To be filled]", "May 2026"]:
        para(doc, item, WD_ALIGN_PARAGRAPH.CENTER, False)
    doc.add_page_break()

    p = para(doc, TITLE, WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    para(doc, "[Student Name]", WD_ALIGN_PARAGRAPH.CENTER, False)
    para(doc, "[College Name], NUIST, Nanjing 210044, China", WD_ALIGN_PARAGRAPH.CENTER, False)
    para(doc, "Abstract: With the development of smart campus construction, campus navigation services have become an important part of digital campus applications. This thesis designs and implements a Smart Campus Navigation Web system based on Spring Boot 3.5.13, Java 21, Maven, MyBatis, MySQL, Thymeleaf, HTML, CSS, native JavaScript, AMap (Gaode Map), QWeather API, and an Intelligent Assistant backend proxy. The system supports POI retrieval, place detail display, route planning, weather display, travel suggestions, Saved Places, Route History Records, and Intelligent Assistant Q&A. The design emphasizes real campus scenarios, clear system boundaries, maintainable module structure, shared data persistence, and practical exception handling.")
    para(doc, "Key Words: Smart Campus; Campus Navigation; AMap; QWeather API; Spring Boot", indent=False)
    doc.add_page_break()

    p = para(doc, "Contents", WD_ALIGN_PARAGRAPH.CENTER, False)
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    for line in [x for x in OUTLINE if x.startswith("Chapter ") or re.match(r"^[1-6]\\.[1-9] ", x)]:
        para(doc, line + " ................................................................", indent=False)
    doc.add_page_break()

    for title in OUTLINE:
        heading(doc, title, level_of(title))
        if title in FIGURES:
            figure(doc, *FIGURES[title])
        if title == "2.1 Development Environment":
            table(doc, "Table 2.1 Development Environment Table", ["Item", "Description"], [
                ("JDK", "Java 21"), ("Backend Framework", "Spring Boot 3.5.13"), ("Build Tool", "Maven"), ("IDE", "IntelliJ IDEA"), ("Database", "MySQL"), ("Frontend", "Thymeleaf, HTML, CSS, Native JavaScript")
            ])
        if title == "2.4 Third-Party Service Technologies":
            table(doc, "Table 2.2 Main Technology Description Table", ["Technology", "Role"], [
                ("AMap JavaScript API", "Map display and frontend map interaction"), ("AMap Web Service", "Route planning service invoked by backend logic"), ("QWeather API", "Current campus weather data provider"), ("Intelligent Assistant Proxy", "Backend proxy invocation for campus scenario Q&A")
            ])
        if title == "3.3.1 Functional Requirement Analysis":
            table(doc, "Table 3.1 Functional Requirement Table", ["ID", "Requirement", "Description"], [
                ("FR-01", "POI Retrieval", "Search campus POIs by keyword and type"), ("FR-02", "Route Planning", "Plan walking or cycling routes by invoking AMap route services"), ("FR-03", "Weather Display", "Display current campus weather through QWeather API"), ("FR-04", "Travel Suggestions", "Generate scene-based campus travel suggestions"), ("FR-05", "Saved Places", "Add, display, rename, and delete shared saved places"), ("FR-06", "Route History Records", "Store, display, reuse, delete, and clear shared route records"), ("FR-07", "Intelligent Assistant", "Use backend proxy invocation for campus scenario Q&A")
            ])
        if title == "3.3.2 Non-Functional Requirement Analysis":
            table(doc, "Table 3.2 Non-Functional Requirement Table", ["Requirement", "Description"], [
                ("Response Speed", "Search and common interface responses should remain smooth under normal conditions"), ("Ease of Use", "Map, sidebar, and route point selection should be clear"), ("Maintainability", "Controller, service, mapper, DTO, VO, and resource layers should remain separated"), ("Security", "External API keys should be encapsulated by backend configuration where applicable"), ("Stability", "External failures and empty results should be handled with controlled prompts")
            ])
        if title == "3.6.3 Description of Main Data Tables":
            table(doc, "Table 3.3 Main Database Table Description", ["Table Name", "Description"], [
                ("poi", "Campus POI master data"), ("saved_place", "Shared Saved Places data"), ("route_history", "Shared Route History Records"), ("covered_path_node", "Sheltered path auxiliary data")
            ])
        if title == "3.8.6 Main System Interface Description Table":
            table(doc, "Table 3.4 Main System Interface Description", ["Interface Path", "Request Method", "Function Description", "Main Parameters", "Return Result"], [
                ("GET /api/v1/pois", "GET", "Query POI list", "name, type, enabled", "List<PoiVO>"), ("GET /api/v1/pois/types", "GET", "Query POI types", "none", "List<String>"), ("GET /api/v1/pois/{id}", "GET", "Get POI detail", "id", "PoiVO"), ("GET /api/v1/routes/walking", "GET", "Plan walking route", "originLng, originLat, destinationLng, destinationLat, optional viaLng/viaLat", "WalkingRouteVO"), ("GET /api/v1/routes/cycling", "GET", "Plan cycling route", "originLng, originLat, destinationLng, destinationLat, optional viaLng/viaLat", "WalkingRouteVO"), ("GET /api/v1/weather/current", "GET", "Get current weather", "none", "WeatherCurrentVO"), ("GET /api/v1/suggestions/context", "GET", "Get context suggestion", "sceneType, poiId, routeDistance, routeDuration, hour", "SuggestionContextVO"), ("GET /api/v1/saved-places", "GET", "List Saved Places", "none", "items, totalCount"), ("POST /api/v1/saved-places", "POST", "Create Saved Place", "SavedPlaceCreateRequest", "SavedPlaceVO"), ("PUT /api/v1/saved-places/{id}/name", "PUT", "Rename Saved Place", "id, name", "SavedPlaceVO"), ("DELETE /api/v1/saved-places/{id}", "DELETE", "Delete Saved Place", "id", "Void"), ("GET /api/v1/route-histories", "GET", "List Route History Records", "query request", "items, totalCount"), ("POST /api/v1/route-histories", "POST", "Create Route History Record", "RouteHistoryCreateRequest", "RouteHistoryVO"), ("DELETE /api/v1/route-histories", "DELETE", "Clear route history", "none", "cleared count"), ("POST /api/v1/assistant/chat", "POST", "Ask Intelligent Assistant", "AssistantChatRequest", "AssistantChatVO")
            ])
        if title == "5.3 Functional Testing":
            table(doc, "Table 5.1 Functional Test Case Table", ["Test ID", "Test Module", "Test Content", "Input Data", "Expected Result", "Actual Result", "Pass or Fail"], [
                ("FT-01", "POI", "Keyword search", "library", "Matched POIs displayed", "To be executed", "Pending"), ("FT-02", "Route", "Walking route", "Start and destination", "Route rendered on map", "To be executed", "Pending"), ("FT-03", "Weather", "Weather display", "Open weather panel", "Weather data displayed", "To be executed", "Pending"), ("FT-04", "Saved Places", "Add saved place", "Selected POI", "Saved place appears", "To be executed", "Pending"), ("FT-05", "Route History Records", "Reuse route record", "Existing route record", "Route points restored", "To be executed", "Pending"), ("FT-06", "Intelligent Assistant", "Campus Q&A", "Campus question", "Answer or fallback shown", "To be executed", "Pending")
            ])
        if title == "5.4 Exception Testing":
            table(doc, "Table 5.2 Exception Test Case Table", ["Test ID", "Test Module", "Test Content", "Input Data", "Expected Result", "Actual Result", "Pass or Fail"], [
                ("ET-01", "Search", "Empty keyword", "empty string", "No crash and empty/default state", "To be executed", "Pending"), ("ET-02", "Search", "Nonexistent place", "unknown place", "No matching result shown", "To be executed", "Pending"), ("ET-03", "Route", "Missing start", "destination only", "Prompt asks for start point", "To be executed", "Pending"), ("ET-04", "Route", "Missing destination", "start only", "Prompt asks for destination", "To be executed", "Pending"), ("ET-05", "Route", "Same start and destination", "same coordinates", "Request rejected", "To be executed", "Pending"), ("ET-06", "Route", "Invalid waypoint", "mismatched via coordinates", "Error prompt returned", "To be executed", "Pending"), ("ET-07", "Weather", "QWeather API failure", "service unavailable", "Fallback prompt shown", "To be executed", "Pending"), ("ET-08", "Assistant", "Assistant failure", "proxy unavailable", "Exception prompt shown", "To be executed", "Pending"), ("ET-09", "Database", "No result", "valid query with no data", "Empty list displayed", "To be executed", "Pending")
            ])
        for text in body_for(title):
            para(doc, text)

    heading(doc, "References", 1)
    refs = [
        "[1] Nanjing University of Information Science and Technology. Undergraduate Thesis Writing Specification and Template[Z]. 2026.",
        "[2] Spring. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/ . Accessed 2026-04-25.",
        "[3] MyBatis. MyBatis 3 User Guide[EB/OL]. https://mybatis.org/mybatis-3/ . Accessed 2026-04-25.",
        "[4] Oracle. Java Platform Documentation[EB/OL]. https://docs.oracle.com/en/java/ . Accessed 2026-04-25.",
        "[5] Apache Maven Project. Maven Documentation[EB/OL]. https://maven.apache.org/guides/ . Accessed 2026-04-25.",
        "[6] MySQL. MySQL Reference Manual[EB/OL]. https://dev.mysql.com/doc/ . Accessed 2026-04-25.",
        "[7] Thymeleaf. Thymeleaf Documentation[EB/OL]. https://www.thymeleaf.org/documentation.html . Accessed 2026-04-25.",
        "[8] AMap Open Platform. JavaScript API Documentation[EB/OL]. https://lbs.amap.com/ . Accessed 2026-04-25.",
        "[9] QWeather. QWeather API Documentation[EB/OL]. https://dev.qweather.com/en/docs/ . Accessed 2026-04-25.",
        "[10] Pressman R S, Maxim B R. Software Engineering: A Practitioner’s Approach[M]. 9th ed. New York: McGraw-Hill Education, 2019.",
    ]
    for r in refs:
        para(doc, r, indent=False)
    heading(doc, "Acknowledgements", 1)
    para(doc, "This section should be completed by the student before final submission. It may thank the supervisor and others who provided guidance, testing assistance, or data support during the graduation design process.")
    heading(doc, "Appendix", 1)
    para(doc, "Appendix materials may include selected API examples, additional test records, or supplementary database initialization notes.")
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
